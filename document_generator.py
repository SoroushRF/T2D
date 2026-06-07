import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fpdf import FPDF, FontFace


def set_cell_background(cell, color_hex):
    """Sets the background color of a table cell in python-docx."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in dxa/twips) for a DOCX cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def export_to_docx(df, output_path, title="Spreadsheet Export"):
    """Exports a pandas DataFrame to a styled Microsoft Word DOCX document."""
    doc = Document()

    # Configure document title
    h1 = doc.add_heading(title, level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h1.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(45, 55, 72)  # Charcoal Slate

    doc.add_paragraph(f"Exported from Spreadsheet Editor. Shape: {df.shape[0]} rows x {df.shape[1]} columns.\n")

    # Add a table
    num_rows = len(df) + 1  # include header
    num_cols = len(df.columns)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Write and style headers
    hdr_cells = table.rows[0].cells
    for col_idx, col_name in enumerate(df.columns):
        cell = hdr_cells[col_idx]
        cell.text = str(col_name)
        set_cell_background(cell, "2D3748")  # Dark Slate Blue
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        # Center header text and change font
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Write and style data rows
    for row_idx, row in df.iterrows():
        row_cells = table.rows[row_idx + 1].cells
        # Zebra striping color
        row_bg = "F7FAFC" if row_idx % 2 == 1 else "FFFFFF"

        for col_idx, val in enumerate(row):
            cell = row_cells[col_idx]

            # Format value
            cell_val = ""
            if not pd.isna(val):
                if isinstance(val, float):
                    cell_val = str(int(val)) if val.is_integer() else f"{val:.4g}"
                else:
                    cell_val = str(val)

            cell.text = cell_val
            set_cell_background(cell, row_bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

            p = cell.paragraphs[0]
            # Left align text, right align numbers
            try:
                # Test if numeric
                float(cell_val)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except ValueError:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(45, 55, 72)

    # Save the file
    doc.save(output_path)


class StyledPDF(FPDF):
    def __init__(self, title_text, filename="", *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.title_text = title_text
        self.filename = filename
        from datetime import datetime

        self.gen_date = datetime.now().strftime("%Y-%m-%d %H:%M")

    def header(self):
        self.set_font("helvetica", "B", 11)
        self.set_text_color(45, 55, 72)
        # Title of the export
        self.cell(0, 8, self.title_text, align="L")

        # Reset x position to print right-aligned text on the same line
        self.set_x(self.l_margin)

        # Right-aligned filename and date
        self.set_font("helvetica", "I", 8)
        self.set_text_color(113, 128, 150)
        file_info = f"File: {self.filename}  |  Generated: {self.gen_date}"
        self.cell(0, 8, file_info, align="R", new_x="LMARGIN", new_y="NEXT")

        self.set_draw_color(226, 232, 240)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", "I", 8)
        self.set_text_color(160, 174, 192)
        total_pages = self.pages_count if hasattr(self, "pages_count") else "{nb}"
        self.cell(0, 10, f"Page {self.page_no()} of {total_pages}", align="C")


def export_to_pdf(
    df,
    output_path,
    title="Spreadsheet Export",
    orientation=None,
    paper_size="A4",
    margin=10,
    filename_template=None,
):
    """Exports a pandas DataFrame to a styled PDF document with layout controls."""
    import os
    from datetime import datetime

    # Handle filename formatting templates
    if filename_template or "{" in output_path:
        template = filename_template or output_path
        now = datetime.now()
        output_path = template.format(
            date=now.strftime("%Y-%m-%d"),
            datetime=now.strftime("%Y%m%d_%H%M%S"),
            title=title.replace(" ", "_"),
        )

    # Determine default page orientation based on column count if not specified
    if not orientation:
        num_cols = len(df.columns)
        orientation = "L" if num_cols > 6 else "P"

    filename = os.path.basename(output_path)
    pdf = StyledPDF(title_text=title, filename=filename, orientation=orientation, unit="mm", format=paper_size)
    pdf.alias_nb_pages()
    pdf.set_margins(left=margin, top=margin, right=margin)
    pdf.add_page()

    # Subtitle with data metadata
    pdf.set_font("helvetica", "I", 9)
    pdf.set_text_color(113, 128, 150)
    pdf.cell(
        0,
        6,
        f"Shape: {df.shape[0]} rows x {df.shape[1]} columns. Generated via T2D Spreadsheet Editor.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(3)

    # Style configuration
    pdf.set_font("helvetica", size=9)

    # Format all cells to strings
    headers = [str(col) for col in df.columns]
    rows = []
    for _, row in df.iterrows():
        row_str = []
        for val in row:
            if pd.isna(val):
                row_str.append("")
            else:
                if isinstance(val, float):
                    row_str.append(str(int(val)) if val.is_integer() else f"{val:.4g}")
                else:
                    row_str.append(str(val))
        rows.append(row_str)

    # Adaptive column width fitting algorithm
    col_max_lens = [len(str(h)) for h in headers]
    for r in rows:
        for i, val in enumerate(r):
            if i < len(col_max_lens):
                col_max_lens[i] = max(col_max_lens[i], len(str(val)))

    total_len = sum(col_max_lens)
    if total_len > 0:
        # Scale to 100% of usable page width
        usable_width = pdf.w - pdf.l_margin - pdf.r_margin
        col_widths = [max(12, int((col_len / total_len) * usable_width)) for col_len in col_max_lens]
    else:
        col_widths = None

    # Draw table using fpdf2 table feature
    with pdf.table(
        col_widths=col_widths,
        text_align="LEFT",
        line_height=6,
        cell_fill_color=(247, 250, 252),
        cell_fill_mode="ROWS",
        headings_style=FontFace(color=(255, 255, 255), fill_color=(45, 55, 72)),
    ) as table:
        # Style headings
        pdf.set_text_color(255, 255, 255)

        # Write headings
        header_row = table.row()
        for h in headers:
            header_row.cell(h)

        # Style body
        pdf.set_text_color(45, 55, 72)

        # Write data rows
        for idx, r in enumerate(rows):
            data_row = table.row()
            for val in r:
                # Right align numbers
                align = "RIGHT"
                try:
                    float(val)
                except ValueError:
                    align = "LEFT"
                data_row.cell(val, align=align)

    pdf.output(output_path)
    return output_path
